#!/usr/bin/env python3
"""Generate Portfolio_AI_Agent_Guide.docx with professional formatting."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_code_block(doc, code, language=""):
    """Add a styled code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Light gray background via shading on paragraph
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F0F0F0",
    })
    pPr.append(shd)


def add_tip(doc, text):
    """Add a tip/note callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "E8F4FD",
    })
    pPr.append(shd)
    run = p.add_run("Tip: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0A, 0x6E, 0xBD)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build():
    doc = Document()

    # -- Default font --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # -- Title --
    title = doc.add_heading("Portfolio AI Agent — Setup Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "A Python script that pulls live market data, sends it to Claude for analysis,\n"
        "and delivers BUY / SELL / HOLD recommendations to your inbox, Slack, Notion, or a local file."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1.  Prerequisites",
        "2.  Install Dependencies",
        "3.  Set Your API Key",
        "4.  Edit Your Holdings",
        "5.  Configure Delivery Channels",
        "6.  Run the Agent",
        "7.  Schedule It Overnight",
        "8.  Agent Flow Diagram",
        "9.  Costs",
        "10. Troubleshooting",
        "11. Customization Ideas",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. Prerequisites
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Prerequisites", level=1)
    doc.add_paragraph(
        "Before you begin, make sure you have the following:"
    )
    items = [
        ("Python 3.10+", "Check with: python3 --version"),
        ("Anthropic API key", "Sign up at console.anthropic.com and create an API key"),
        ("Internet connection", "The script fetches live prices from Yahoo Finance"),
    ]
    for bold_part, rest in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(f" — {rest}")

    # ═══════════════════════════════════════════════════════════════════
    # 2. Install Dependencies
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Install Dependencies", level=1)
    doc.add_paragraph("The agent uses only two external packages:")
    add_code_block(doc, "pip install yfinance anthropic")
    doc.add_paragraph(
        "That's it — everything else comes from the Python standard library. "
        "If you use a virtual environment (recommended):"
    )
    add_code_block(doc, "python3 -m venv .venv\nsource .venv/bin/activate      # Mac/Linux\n.venv\\Scripts\\activate         # Windows\npip install yfinance anthropic")

    # ═══════════════════════════════════════════════════════════════════
    # 3. Set Your API Key
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Set Your API Key", level=1)

    doc.add_heading("Mac / Linux", level=2)
    add_code_block(doc, 'export ANTHROPIC_API_KEY="sk-ant-api03-..."')
    doc.add_paragraph(
        "To make it permanent, add that line to ~/.zshrc (Mac) or ~/.bashrc (Linux), "
        "then run: source ~/.zshrc"
    )

    doc.add_heading("Windows (PowerShell)", level=2)
    add_code_block(doc, '$env:ANTHROPIC_API_KEY = "sk-ant-api03-..."')
    doc.add_paragraph(
        "To persist it: System Properties → Environment Variables → New User Variable. "
        "Name: ANTHROPIC_API_KEY, Value: your key."
    )

    # ═══════════════════════════════════════════════════════════════════
    # 4. Edit Your Holdings
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Edit Your Holdings", level=1)
    doc.add_paragraph(
        "Open portfolio_agent.py and update the PORTFOLIO dict at the top with your actual positions:"
    )
    add_code_block(doc,
        'PORTFOLIO = {\n'
        '    "AAPL": {"shares": 50, "avg_cost": 142.00, "account": "Taxable"},\n'
        '    "VTI":  {"shares": 100, "avg_cost": 210.50, "account": "Roth IRA"},\n'
        '    # Add your positions here...\n'
        '}'
    )

    doc.add_paragraph("Each entry accepts the following fields:")

    # Fields table
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Field", "Type", "Description"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    rows_data = [
        ("shares", "float", "Number of shares you own"),
        ("avg_cost", "float", "Your average cost per share"),
        ("account", "string", "Where it's held (e.g. \"Roth IRA\", \"Taxable\")"),
        ("yf_ticker", "string (optional)", "Override if Yahoo Finance uses a different symbol"),
        ("skip_analysis", "bool (optional)", "Set True for money market / cash positions"),
    ]
    for r, (field, typ, desc) in enumerate(rows_data, start=1):
        table.rows[r].cells[0].text = field
        table.rows[r].cells[1].text = typ
        table.rows[r].cells[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph(
        "Also update the RECURRING list if you have regular DCA (dollar-cost averaging) "
        "transfers — Claude uses this context to evaluate your allocation strategy."
    )

    # ═══════════════════════════════════════════════════════════════════
    # 5. Configure Delivery
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Configure Delivery Channels", level=1)

    # Local
    doc.add_heading("Local File (default — always on)", level=2)
    doc.add_paragraph(
        "Reports automatically save to a reports/ folder next to the script as dated Markdown files. "
        "No configuration needed."
    )

    # Email
    doc.add_heading("Email (Gmail example)", level=2)
    p = doc.add_paragraph(style="List Number")
    p.add_run("Enable 2-factor authentication on your Google account")
    p = doc.add_paragraph(style="List Number")
    p.add_run("Generate an App Password at myaccount.google.com/apppasswords")
    p = doc.add_paragraph(style="List Number")
    p.add_run("Fill in the config block in portfolio_agent.py:")

    add_code_block(doc,
        'SEND_EMAIL = True\n'
        'EMAIL_FROM = "you@gmail.com"\n'
        'EMAIL_TO   = "you@gmail.com"\n'
        'SMTP_HOST  = "smtp.gmail.com"\n'
        'SMTP_PORT  = 587\n'
        'SMTP_USER  = "you@gmail.com"\n'
        'SMTP_PASS  = "xxxx xxxx xxxx xxxx"  # 16-char app password'
    )

    add_tip(doc, "Use an App Password, never your real Gmail password. The script uses STARTTLS for encryption.")

    # Slack
    doc.add_heading("Slack", level=2)
    p = doc.add_paragraph(style="List Number")
    p.add_run("Create an Incoming Webhook in your Slack workspace (api.slack.com/messaging/webhooks)")
    p = doc.add_paragraph(style="List Number")
    p.add_run("Copy the webhook URL into the config:")

    add_code_block(doc,
        'SEND_SLACK        = True\n'
        'SLACK_WEBHOOK_URL = "https://hooks.slack.com/services/T.../B.../xxx"'
    )

    # Notion
    doc.add_heading("Notion", level=2)
    p = doc.add_paragraph(style="List Number")
    p.add_run("Create an internal integration at notion.so/my-integrations")
    p = doc.add_paragraph(style="List Number")
    p.add_run("Share a database with the integration (the database needs a Name title property and a Date date property)")
    p = doc.add_paragraph(style="List Number")
    p.add_run("Add credentials to the config:")

    add_code_block(doc,
        'POST_NOTION        = True\n'
        'NOTION_API_KEY     = "secret_..."\n'
        'NOTION_DATABASE_ID = "abc123..."'
    )

    # ═══════════════════════════════════════════════════════════════════
    # 6. Run It
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Run the Agent", level=1)
    add_code_block(doc, "python3 portfolio_agent.py")
    doc.add_paragraph("You'll see output like this:")
    add_code_block(doc,
        '═══ Portfolio Agent — 2026-02-26 ═══\n\n'
        'Fetching market data...\n'
        '  FXAIX... $221.45\n'
        '  QQQM...  $258.12\n'
        '  SPY...   $612.30\n'
        '  ...\n\n'
        '  Total value:     $25,432.10\n'
        '  Total cost:      $22,891.55\n'
        '  Gain/Loss:       +$2,540.55 (+11.10%)\n\n'
        'Sending to Claude for analysis...\n'
        '────────────────────────────────\n'
        '## Portfolio Overview\n'
        '...\n'
        '## Per-Position Analysis\n'
        '**FXAIX — BUY** ...\n'
        '────────────────────────────────\n\n'
        'Delivering report...\n'
        '  Saved → reports/portfolio_report_2026-02-26.md\n\n'
        'Done.'
    )

    # ═══════════════════════════════════════════════════════════════════
    # 7. Schedule It Overnight
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Schedule It Overnight", level=1)
    doc.add_paragraph(
        "Run the agent automatically every weekday morning before market open "
        "so you have fresh analysis waiting for you."
    )

    # Mac
    doc.add_heading("Mac (cron)", level=2)
    doc.add_paragraph("Open your crontab:")
    add_code_block(doc, "crontab -e")
    doc.add_paragraph("Add this line to run at 6:30 AM, Monday through Friday:")
    add_code_block(doc,
        '30 6 * * 1-5  /usr/bin/env ANTHROPIC_API_KEY="sk-ant-..." '
        '/usr/local/bin/python3 /Users/you/wealthpilot/portfolio_agent.py '
        '>> ~/portfolio_log.txt 2>&1'
    )
    add_tip(doc,
        "cron doesn't source your shell profile, so you must either inline the "
        "ANTHROPIC_API_KEY (as shown above) or wrap the command in a shell script "
        "that sources ~/.zshrc first."
    )

    # Windows
    doc.add_heading("Windows (Task Scheduler)", level=2)
    steps = [
        "Open Task Scheduler → Create Basic Task",
        "Trigger: Weekly, check Mon–Fri, set time to 6:30 AM",
        "Action: Start a Program",
        "  Program: python",
        "  Arguments: C:\\path\\to\\portfolio_agent.py",
        "  Start in: C:\\path\\to\\",
        'In the task properties, check "Run whether user is logged on or not"',
        "Add ANTHROPIC_API_KEY as a system environment variable (System Properties → Environment Variables)",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.space_after = Pt(2)

    # Cloud
    doc.add_heading("Cloud (always-on, no laptop needed)", level=2)

    doc.add_heading("Option A: GitHub Actions (free tier — 2,000 min/month)", level=3)
    doc.add_paragraph("Create .github/workflows/portfolio.yml in your repo:")
    add_code_block(doc,
        'name: Portfolio Agent\n'
        'on:\n'
        '  schedule:\n'
        '    - cron: "30 10 * * 1-5"   # 10:30 UTC = 6:30 AM ET\n'
        '  workflow_dispatch:            # manual trigger button\n\n'
        'jobs:\n'
        '  analyze:\n'
        '    runs-on: ubuntu-latest\n'
        '    steps:\n'
        '      - uses: actions/checkout@v4\n'
        '      - uses: actions/setup-python@v5\n'
        '        with:\n'
        '          python-version: "3.12"\n'
        '      - run: pip install yfinance anthropic\n'
        '      - run: python portfolio_agent.py\n'
        '        env:\n'
        '          ANTHROPIC_API_KEY: ${{ secrets.ANTHROPIC_API_KEY }}'
    )
    doc.add_paragraph(
        "Store your ANTHROPIC_API_KEY in the repo's Settings → Secrets and variables → Actions → New repository secret."
    )

    doc.add_heading("Option B: Other cloud platforms", level=3)
    other_cloud = [
        ("AWS Lambda + EventBridge", "Serverless, pay-per-invocation (~$0/month for this use case)"),
        ("Google Cloud Functions + Cloud Scheduler", "Similar to AWS, generous free tier"),
        ("$5/month VPS (DigitalOcean, Hetzner, etc.)", "Just install Python, clone repo, set up cron"),
    ]
    for name, desc in other_cloud:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name)
        run.bold = True
        p.add_run(f" — {desc}")

    # ═══════════════════════════════════════════════════════════════════
    # 8. Agent Flow
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Agent Flow Diagram", level=1)
    doc.add_paragraph("The agent runs through four stages in sequence:")

    # Flow as a table for clean rendering in Word
    flow_table = doc.add_table(rows=2, cols=4)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stage_headers = ["1. Fetch Data", "2. Calculate Metrics", "3. Claude Analysis", "4. Deliver Report"]
    stage_details = [
        "Yahoo Finance:\n• Live prices\n• 52-week hi/lo\n• P/E ratios\n• Analyst targets",
        "Per position:\n• Portfolio weight %\n• 1-month return\n• 3-month return\n• Gain/loss vs cost",
        "Claude API:\n• BUY/SELL/HOLD\n  per position\n• Rationale\n• Action items\n• DCA review",
        "Channels:\n• Local .md file\n• Email (SMTP)\n• Slack webhook\n• Notion page",
    ]
    for i in range(4):
        cell = flow_table.rows[0].cells[i]
        cell.text = stage_headers[i]
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

        cell = flow_table.rows[1].cells[i]
        cell.text = stage_details[i]
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

    # ═══════════════════════════════════════════════════════════════════
    # 9. Costs
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("9. Costs", level=1)

    cost_table = doc.add_table(rows=4, cols=3)
    cost_table.style = "Light Grid Accent 1"
    cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cost_headers = ["Component", "Cost", "Notes"]
    for i, h in enumerate(cost_headers):
        cell = cost_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    cost_rows = [
        ("Yahoo Finance", "Free", "yfinance scrapes public data, no API key needed"),
        ("Claude API (Sonnet)", "~$0.01–0.05 / run", "Recommended for daily use"),
        ("Claude API (Opus)", "~$0.10–0.30 / run", "Deeper analysis, use for weekly deep-dives"),
    ]
    for r, (comp, cost, note) in enumerate(cost_rows, start=1):
        cost_table.rows[r].cells[0].text = comp
        cost_table.rows[r].cells[1].text = cost
        cost_table.rows[r].cells[2].text = note

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Monthly estimate: ")
    run.bold = True
    p.add_run("$0.50–$6/month running once per weekday, depending on model choice.")

    # ═══════════════════════════════════════════════════════════════════
    # 10. Troubleshooting
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("10. Troubleshooting", level=1)

    trouble_table = doc.add_table(rows=7, cols=2)
    trouble_table.style = "Light Grid Accent 1"
    trouble_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    trouble_table.rows[0].cells[0].text = "Problem"
    trouble_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    trouble_table.rows[0].cells[1].text = "Fix"
    trouble_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    troubles = [
        ("ModuleNotFoundError: yfinance", "pip install yfinance"),
        ("ModuleNotFoundError: anthropic", "pip install anthropic"),
        ("Ticker returns no price", "Check the symbol on finance.yahoo.com — mutual funds may use a different ticker. Use the yf_ticker field to override."),
        ("Email authentication fails", "Make sure you're using a Gmail App Password (16 characters), not your account password."),
        ("Cron doesn't run", "cron doesn't load your shell profile. Set ANTHROPIC_API_KEY inline in the crontab entry."),
        ("Rate limited by Yahoo Finance", "Add time.sleep(1) between fetches in the loop, or reduce run frequency."),
    ]
    for r, (prob, fix) in enumerate(troubles, start=1):
        trouble_table.rows[r].cells[0].text = prob
        trouble_table.rows[r].cells[1].text = fix

    # ═══════════════════════════════════════════════════════════════════
    # 11. Customization Ideas
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("11. Customization Ideas", level=1)

    ideas = [
        ("Change the Claude model", 'Set CLAUDE_MODEL = "claude-opus-4-6" for deeper, more nuanced analysis'),
        ("Add more tickers", "Just add entries to the PORTFOLIO dict — no other changes needed"),
        ("Track crypto spot prices", 'Use "yf_ticker": "BTC-USD" for the live Bitcoin price instead of an ETF'),
        ("Compare to benchmarks", "Add SPY/QQQ performance to the prompt context for relative analysis"),
        ("Historical tracking", "Reports save with dates so you can diff them over time and track trends"),
        ("Multiple portfolios", "Duplicate the PORTFOLIO dict or create a config file to run the agent for different accounts"),
    ]
    for bold_part, desc in ideas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(f" — {desc}")

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WealthPilot — Portfolio AI Agent")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──
    out_path = "/Users/dany.zisman/wealthpilot/Portfolio_AI_Agent_Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
